import os
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from collections import defaultdict

def clean_email_text(text):
    """Clean markdown and other formatting from text before email extraction"""
    if not text:
        return ""
    
    # Remove markdown link formatting: [email](mailto:email) -> email
    text = re.sub(r'\[([^\]]+?)\]\(mailto:([^\)]+?)\)', r'\2', text)
    
    # Remove remaining markdown links: [text](url) -> text  
    text = re.sub(r'\[([^\]]+?)\]\([^\)]+?\)', r'\1', text)
    
    # Remove other markdown formatting
    text = re.sub(r'\*\*([^*]+?)\*\*', r'\1', text)  # **bold**
    text = re.sub(r'\*([^*]+?)\*', r'\1', text)        # *italic*
    
    return text

def extract_emails_from_text(text, debug=False):
    """Extract email addresses using comprehensive patterns and advanced cleaning"""
    if not text:
        return []
    
    # Convert to string and normalize whitespace
    text = str(text).replace('\r\n', '\n').replace('\r', '\n')
    
    # Clean all markdown formatting first
    text = clean_email_text(text)
    
    # Multiple comprehensive email patterns
    patterns = [
        # Standard comprehensive pattern
        r'\b[A-Za-z0-9]([A-Za-z0-9._%-]*[A-Za-z0-9])?@[A-Za-z0-9]([A-Za-z0-9.-]*[A-Za-z0-9])?\.[A-Za-z]{2,6}\b',
        
        # Pattern for emails with underscores and dots
        r'[A-Za-z0-9][A-Za-z0-9._%-]*@[A-Za-z0-9][A-Za-z0-9.-]*\.[A-Za-z]{2,6}',
        
        # Pattern for emails with spaces (PDF artifacts)
        r'[A-Za-z0-9][A-Za-z0-9._%-]*\s*@\s*[A-Za-z0-9][A-Za-z0-9.-]*\s*\.\s*[A-Za-z]{2,6}',
        
        # Pattern for emails with newlines or weird spacing
        r'[A-Za-z0-9._%-]+\s*@\s*[A-Za-z0-9.-]+\s*\.\s*[A-Za-z]{2,6}',
        
        # Very permissive pattern for malformed emails
        r'[A-Za-z0-9._%-]+[@＠][A-Za-z0-9.-]+[\.|．][A-Za-z]{2,6}',
        
        # Pattern to catch emails split across lines
        r'[A-Za-z0-9._%-]+\n?@\n?[A-Za-z0-9.-]+\n?\.\n?[A-Za-z]{2,6}'
    ]
    
    all_emails = []
    
    if debug:
        print(f"    Text length: {len(text)} characters")
        print(f"    @ symbols found: {text.count('@')}")
    
    # Apply each pattern
    for i, pattern in enumerate(patterns):
        try:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if matches:
                all_emails.extend(matches)
                if debug:
                    print(f"    Pattern {i+1}: found {len(matches)} matches")
        except Exception as e:
            if debug:
                print(f"    Pattern {i+1}: error - {e}")
            continue
    
    # Manual extraction around @ symbols
    lines = text.split('\n')
    for line in lines:
        if '@' in line:
            # Split line by common delimiters
            words = re.split(r'[\s,;:|()[\]{}\"\'`<>]+', line)
            for word in words:
                if '@' in word and len(word) > 4:
                    # Clean potential email
                    cleaned = re.sub(r'[^a-zA-Z0-9@._-]', '', word)
                    if '@' in cleaned and '.' in cleaned:
                        all_emails.append(cleaned)
    
    # Advanced cleaning and validation
    cleaned_emails = []
    for email in all_emails:
        try:
            # Convert to string and initial cleaning
            email = str(email).strip()
            
            # Additional markdown cleaning
            email = clean_email_text(email)
            
            # Remove all types of whitespace and line breaks
            email = re.sub(r'[\s\n\r\t]', '', email)
            
            # Final validation
            if not email or email.count('@') != 1:
                continue
            
            # Check for valid email format
            if not re.match(r'^[A-Za-z0-9][A-Za-z0-9._%-]*@[A-Za-z0-9][A-Za-z0-9.-]*\.[A-Za-z]{2,6}$', email):
                continue
            
            # Additional security checks
            if '.' not in email.split('@')[1]:
                continue
            
            cleaned_emails.append(email)
        
        except Exception as e:
            if debug:
                print(f"    Error cleaning email: {e}")
            continue
    
    return cleaned_emails

def extract_from_pdf(file_path, debug=False):
    """Extract text from PDF files"""
    try:
        import pdfplumber
        
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                text += "\n"
        
        if debug:
            print(f"    PDF: {len(text)} chars, {text.count('@')} @ symbols")
        
        return text
    except Exception as e:
        if debug:
            print(f"    PDF error: {e}")
        return ""

def extract_from_docx(file_path, debug=False):
    """Extract text from DOCX files"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        if debug:
            print(f"    DOCX: {len(text)} chars, {text.count('@')} @ symbols")
        
        return text
    except Exception as e:
        if debug:
            print(f"    DOCX error: {e}")
        return ""

def extract_from_doc(file_path, debug=False):
    """Extract text from DOC files"""
    try:
        import win32com.client
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        doc = word.Documents.Open(str(file_path.absolute()))
        
        # Get main document text
        text = doc.Content.Text
        
        # Try to get text from headers/footers
        for section in doc.Sections:
            try:
                if section.Headers:
                    for header in section.Headers:
                        text += header.Range.Text + "\n"
                if section.Footers:
                    for footer in section.Footers:
                        text += footer.Range.Text + "\n"
            except:
                continue
        
        doc.Close(SaveChanges=False)
        word.Quit()
        
        if debug:
            print(f"    DOC: {len(text)} chars, {text.count('@')} @ symbols")
        
        return text
    except Exception as e:
        if debug:
            print(f"    DOC error: {e}")
        return ""

def extract_emails_from_file(file_path, debug=False):
    """Extract emails from a single file"""
    file_path = Path(file_path)
    file_type = file_path.suffix.lower()
    
    text = ""
    
    try:
        if file_type == '.pdf':
            text = extract_from_pdf(file_path, debug)
        elif file_type == '.docx':
            text = extract_from_docx(file_path, debug)
        elif file_type == '.doc':
            text = extract_from_doc(file_path, debug)
        else:
            return []
    except Exception as e:
        if debug:
            print(f"Error extracting from {file_path.name}: {e}")
        return []
    
    if not text or len(text.strip()) < 20:
        return []
    
    emails = extract_emails_from_text(text, debug)
    return emails

def extract_emails_from_uploaded_files(file_paths):
    """
    Extract emails from multiple uploaded files
    Returns: (emails_list, stats_dict, file_mapping_list)
    """
    all_emails = []
    file_mapping = []
    stats = {
        'processed': 0,
        'with_text': 0,
        'with_at_symbols': 0,
        'with_emails': 0,
        'pdf_processed': 0,
        'docx_processed': 0,
        'doc_processed': 0,
        'pdf_emails': 0,
        'docx_emails': 0,
        'doc_emails': 0
    }
    
    for file_path in file_paths:
        file_path = Path(file_path)
        file_type = file_path.suffix.lower()
        
        # Handle ZIP files - extract and process contents
        if file_type == '.zip':
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    # Extract to temp folder
                    temp_extract_dir = Path(file_path).parent / f"temp_zip_{file_path.stem}"
                    zip_ref.extractall(temp_extract_dir)
                    
                    # Process extracted files
                    for extracted_file in temp_extract_dir.rglob('*'):
                        if extracted_file.suffix.lower() in {'.pdf', '.docx', '.doc'}:
                            emails = extract_emails_from_file(extracted_file, debug=False)
                            if emails:
                                all_emails.extend(emails)
                                for email in emails:
                                    file_mapping.append({
                                        'email': email,
                                        'filename': extracted_file.name,
                                        'file_type': extracted_file.suffix.lower()
                                    })
                                stats['with_emails'] += 1
                    
                    # Cleanup temp folder
                    import shutil
                    shutil.rmtree(temp_extract_dir, ignore_errors=True)
            except Exception as e:
                print(f"Error processing ZIP file: {e}")
        else:
            # Process individual files
            text = ""
            try:
                if file_type == '.pdf':
                    text = extract_from_pdf(file_path, debug=False)
                    stats['pdf_processed'] += 1
                elif file_type == '.docx':
                    text = extract_from_docx(file_path, debug=False)
                    stats['docx_processed'] += 1
                elif file_type == '.doc':
                    text = extract_from_doc(file_path, debug=False)
                    stats['doc_processed'] += 1
                
                stats['processed'] += 1
            except Exception as e:
                print(f"Error extracting from {file_path.name}: {e}")
                continue
            
            if not text or len(text.strip()) < 20:
                continue
            
            stats['with_text'] += 1
            
            # Check for @ symbols
            at_count = text.count('@')
            if at_count > 0:
                stats['with_at_symbols'] += 1
            
            # Extract emails
            emails = extract_emails_from_text(text, debug=False)
            
            if emails:
                stats['with_emails'] += 1
                all_emails.extend(emails)
                
                # Track file type stats
                if file_type == '.pdf':
                    stats['pdf_emails'] += 1
                elif file_type == '.docx':
                    stats['docx_emails'] += 1
                elif file_type == '.doc':
                    stats['doc_emails'] += 1
                
                # Track file mapping
                for email in emails:
                    file_mapping.append({
                        'email': email,
                        'filename': file_path.name,
                        'file_type': file_type
                    })
    
    return all_emails, stats, file_mapping
